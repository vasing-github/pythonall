# -*- coding: utf-8 -*-
import os

from datetime import datetime
from bazhouqu.kaipumodify.cfg import conf
import pytz
import requests
import xlrd
from openpyxl import load_workbook
from xlutils.copy import copy
from docx import Document
import pypandoc
import subprocess
import fitz
import sys
from pathlib import Path


def upload_new_file(jid, bz_gov_id, filename, columnId):
    cookies = {
        'authenticatecenterjsessionid': jid,
        'bz_govc_SHIROJSESSIONID': bz_gov_id,
    }

    headers = {
        'Accept': '*/*',
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6',
        'Connection': 'keep-alive',
        # 'Content-Type': 'multipart/form-data; boundary=----WebKitFormBoundaryi3TLT1PsLfm0Pa9i',
        'Origin': f'http://10.15.3.133:{conf.jiyuehua_port}',
        'Referer': f'http://10.15.3.133:{conf.jiyuehua_port}/ewebeditor/ewebeditor.htm?id=content&instanceid=content&style=Lstandard3&cusdir=/{conf.jiyuehua_siteid}/{columnId}&savepathfilename=d_savepathfilename&titleimage=1&extcss=/pageStyle/getCssByColumn/{columnId}&fixwidth=1078px',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/134.0.0.0 Safari/537.36 Edg/134.0.0.0'
    }

    url = f'http://10.15.3.133:{conf.jiyuehua_port}/ewebeditor/jsp/upload.jsp?style=Lstandard3&cusdir=/{conf.jiyuehua_siteid}/{columnId}&skey=&h=10.15.3.133&o=http://10.15.3.133:{conf.jiyuehua_port}&action=mfu&type=file&blockflag=end'
    script_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(script_dir, 'downloadfile', filename)
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"待上传文件不存在: {file_path}")

    try:
        with open(file_path, 'rb') as f:
            files = {
                'uploadfile': (filename, f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'),
                'originalfilename': (None, filename)
            }
            response = requests.post(url, cookies=cookies, headers=headers, files=files, verify=False)
    except IOError as e:
        raise RuntimeError(f"无法读取文件 {file_path}: {str(e)}")

    print(f"响应状态码: {response.status_code}")

    print(response.text)
    return response.json()['url']


def download_file(url, local_filename):
    response = requests.get(url, stream=True, verify=False)
    response.raise_for_status()

    script_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(script_dir, 'downloadfile',local_filename)

    with open(file_path, 'wb') as file:
        for chunk in response.iter_content(chunk_size=8192):
            file.write(chunk)

    print(f"文件已下载并保存为: {file_path}")
    # time.sleep(5)
    return file_path



def modify_file_xls(file_path, unique_replacements_list):
    # 打开xls文件
    workbook_xls = xlrd.open_workbook(file_path, formatting_info=True)
    workbook_copy = copy(workbook_xls)


    # 遍历所有的sheet
    for sheet_index in range(workbook_xls.nsheets):
        sheet_xls = workbook_xls.sheet_by_index(sheet_index)
        sheet_copy = workbook_copy.get_sheet(sheet_index)

        # 遍历所有的单元格
        for row in range(sheet_xls.nrows):
            for col in range(sheet_xls.ncols):
                cell_value = sheet_xls.cell_value(row, col)

                for i in unique_replacements_list:

                    if isinstance(cell_value, str) and i["sensitiveWords"] in cell_value:
                        new_value = cell_value.replace(i["sensitiveWords"], i["recommendUpdate"])
                        sheet_copy.write(row, col, new_value)
                        print(
                            f"Sheet '{sheet_xls.name}' 中的单元格 '{row + 1},{col + 1}' 的值已从 '{cell_value}' 替换为 '{new_value}'")

    # 保存修改后的xls文件
    workbook_copy.save(file_path)
    print(f"文件 '{file_path}' 已成功修改并保存")


def modify_file_xlsx(file_path, unique_replacements_list):
    # 加载工作簿
    workbook = load_workbook(file_path)

    # 遍历所有的 sheet
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]

        # 遍历所有的单元格
        for row in sheet.iter_rows():
            for cell in row:
                for i in unique_replacements_list:
                    if isinstance(cell.value, str) and i["sensitiveWords"] in cell.value:
                        new_value = cell.value.replace(i["sensitiveWords"], i["recommendUpdate"])
                        cell.value = new_value
                        print(f"Sheet '{sheet_name}' 中的单元格 '{cell.coordinate}' 的值已从 '{i['sensitiveWords']}' 替换为 '{i['recommendUpdate']}'")

    # 保存修改后的工作簿
    workbook.save(file_path)
    print(f"文件 '{file_path}' 已成功修改并保存")


def modify_file_docx(file_path, unique_replacements_list):
    # 加载文档
    doc = Document(file_path)

    # 遍历所有段落
    for para in doc.paragraphs:
        for i in unique_replacements_list:
            if i["sensitiveWords"] in para.text:
                new_text = para.text.replace(i["sensitiveWords"], i["recommendUpdate"])
                para.text = new_text
                print(f"段落中的文本已从 '{i['sensitiveWords']}' 替换为 '{i['recommendUpdate']}'")

    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for i in unique_replacements_list:
                    if i["sensitiveWords"] in cell.text:
                        new_text = cell.text.replace(i["sensitiveWords"], i["recommendUpdate"])
                        cell.text = new_text
                        print(f"表格中的单元格文本已从 '{i['sensitiveWords']}' 替换为 '{i['recommendUpdate']}'")

    # 保存修改后的文档
    doc.save(file_path)
    print(f"文件 '{file_path}' 已成功修改并保存")




def modify_file(filename, item):
    script_dir = os.path.dirname(os.path.abspath(__file__))

    file_path = os.path.join(script_dir, 'downloadfile',filename)

    _, file_extension = os.path.splitext(file_path)

    unique_replacements = {}
    for i in item:
        key = (i['sensitiveWords'], i['recommendUpdate'])
        if key not in unique_replacements:
            unique_replacements[key] = i

    # 将字典转换为列表
    unique_replacements_list = list(unique_replacements.values())

    if file_extension == '.xls':
        modify_file_xls(file_path, unique_replacements_list)
    elif file_extension == '.xlsx':
        modify_file_xlsx(file_path, unique_replacements_list)
    elif  file_extension == '.docx':
        modify_file_docx(file_path, unique_replacements_list)
    elif file_extension == '.doc' :
        modify_file_doc(file_path, unique_replacements_list)
        return 1
    elif file_extension == '.pdf' or file_extension == 'PDF':
        modify_file_pdf(file_path, unique_replacements_list)
    return 0


def generate_splits(old_text):
    """生成所有两段式拆分组合（包含完整未拆分情况）"""
    splits = [(old_text[:i], old_text[i:]) for i in range(len(old_text), 0, -1)]
    splits.append((old_text, ""))  # 包含完整文本情况
    return splits


def find_crossline_matches(page, old_text):
    text_blocks = [b for b in page.get_text("blocks") if b[6] == 0]
    matches = []

    # 遍历所有可能的拆分组合（优先最长前缀）
    for prefix, suffix in generate_splits(old_text):
        for i in range(len(text_blocks) - 1):
            curr_block = text_blocks[i][4].replace('\n', '')
            next_block = text_blocks[i + 1][4].replace('\n', '')

            # 当前块必须严格以prefix结尾
            if not curr_block.endswith(prefix):
                continue

            # 下一个块必须严格以suffix开头
            if not next_block.startswith(suffix):
                continue

            # 计算合并区域
            rect_curr = fitz.Rect(text_blocks[i][:4])
            rect_next = fitz.Rect(text_blocks[i + 1][:4])
            combined_rect = rect_curr | rect_next

            matches.append({
                "rects": [rect_curr, rect_next],  # 保留原始坐标信息
                "combined": combined_rect,  # 合并后的区域坐标
                "original_split": (prefix, suffix),  # 原始文本拆分方式
                # "prev_block_text": text_blocks[i][4],  # 前一个块的原始文本（含换行符）
                # "next_block_text": text_blocks[i + 1][4],  # 后一个块的原始文本
                "clean_prev": curr_block,  # 前一个块处理后的文本（已去换行）
                "clean_next": next_block  # 后一个块处理后的文本
            })
            break  # 优先处理最长匹配

    return matches


def smart_replace(page, old_text, new_text):

    # 单行替换
    single_matches = page.search_for(old_text)

    for rect in single_matches:
        print(f"  🖍️ 正在替换单行文本（）")
        print(f"   → 位置：{rect} | 原文本：'{old_text}' → 新文本：'{new_text}'")
        print(f"   → 字体：china-s 11pt 颜色：黑色")
        page.add_redact_annot(rect)
        page.apply_redactions()
        page.insert_text(rect.bl + (0, -5), new_text,
                         fontsize=11,
                         fontname="china-s",
                         color=(0, 0, 0))

        # 跨行替换
    cross_matches = find_crossline_matches(page, old_text)

    for match in cross_matches:
        print(f"\n  🌐 处理跨行匹配")
        print(f"   → 拆分方式：'{match['original_split'][0]}' + '{match['original_split'][1]}'")
        print(f"   → 合并区域：{match['combined']}")

        # 显示实际修改内容
        final_text = f"{match['clean_prev'][:-len(match['original_split'][0])]}{new_text}{match['clean_next'][len(match['original_split'][1]):]}"
        print(f"   → 生成文本：'{final_text[:30]}...'")

        # 删除原区域提示
        print(f"   🗑️ 删除原始区域：")
        # 删除原始区域
        for rect in match["rects"]:
            page.add_redact_annot(rect)
        page.apply_redactions()

        # 构建完整文本内容
        prefix, suffix = match["original_split"]
        prev_remain = match["clean_prev"][:-len(prefix)] if match["clean_prev"].endswith(prefix) else match[
            "clean_prev"]
        next_remain = match["clean_next"][len(suffix):] if match["clean_next"].startswith(suffix) else match[
            "clean_next"]

        # 构建最终完整文本（前内容+新内容+后内容）
        final_text = f"{prev_remain}{new_text}{next_remain}"

        # 将完整文本平均分割为两行（精确居中分割）
        split_index = len(final_text) // 2  # 整体分割点
        line1_text = final_text[:split_index]
        line2_text = final_text[split_index:]

        # 计算插入位置（基于合并区域的垂直居中基准点）
        combined_rect = match["combined"]
        font_size = 10
        line_spacing = font_size * 1.5  # 行间距

        # 基准点计算（在合并区域垂直居中位置）
        base_y = combined_rect.y0 + (combined_rect.height - line_spacing) / 2  # 两行整体垂直居中
        base_point = (combined_rect.x0, base_y)  # 左对齐基准点

        # 两行位置计算（PyMuPDF坐标系y轴向下增长）
        line1_pos = base_point
        line2_pos = (base_point[0], base_point[1] + line_spacing)
        print(f"   ✏️ 插入新文本（字体：china-s 10pt 颜色：黑色）")
        print(f"    行1位置：{match['combined'].x0, match['combined'].y0 + 10}")
        print(f"    行2位置：{match['combined'].x0, match['combined'].y0 + 25}")
        # 插入两行文本（确保顺序正确）
        page.insert_text(
            line1_pos,
            line1_text,
            fontsize=font_size,
            color=(0, 0, 0),
            fontname="china-s",
            overlay=True
        )

        page.insert_text(
            line2_pos,
            line2_text,
            fontsize=font_size,
            color=(0, 0, 0),
            fontname="china-s",
            overlay=True
        )

def modify_file_pdf(file_path, unique_replacements_list):
    doc = fitz.open(file_path)
    for page in doc:
        for replacement in unique_replacements_list:
            sensitive_word = replacement["sensitiveWords"]
            new_word = replacement["recommendUpdate"]
            smart_replace(page,sensitive_word,new_word)
    temp_file = file_path.replace(".pdf", "_temp.pdf")
    doc.save(temp_file, garbage=4, deflate=True)
    doc.close()

    # 用临时文件替换原文件
    os.replace(temp_file, file_path)
    print("\n✅ 所有页面处理完成")
    print(f"💾 正在保存修改到: {file_path}")

def modify_file_doc(file_path_old, unique_replacements_list):
    try:
        output = doc_to_docx(file_path_old)
        print(f"转换成功！文件保存至：{output}")
        print(f"文件大小：{os.path.getsize(output) / 1024:.2f} KB")
    except Exception as e:
        print(f"错误：{str(e)}")
    # file_path_new = file_path_old.replace('.doc', '.docx')
    # pypandoc.convert_file(file_path_old, 'docx', outputfile=file_path)
    modify_file_docx(output, unique_replacements_list)


def uploadfile(jid, bz_gov_id, file_name, path_excel, parent_t, article_t, content_id):
    cookies = {
        'historyCookie': '%E5%8E%BF%E4%BA%BA%E6%B0%91%E6%94%BF%E5%BA%9C%E5%8A%9E%E5%85%AC%E5%AE%A4%2C%E4%B8%AD%E5%8D%8E%E4%BA%BA%E6%B0%91%E5%85%B1%E5%92%8C%E5%9B%BD%E6%95%99%E8%82%B2%E9%83%A8%2C%E4%BA%BA%E6%B0%91%E6%97%A5%E6%8A%A5%2C%E5%8E%BF%E8%87%AA%E7%84%B6%E8%B5%84%E6%BA%90%E5%92%8C%E8%A7%84%E5%88%92%E5%B1%80%2C%E5%9B%9B%E5%B7%9D%E7%9C%81%E6%95%99%E8%82%B2%E8%80%83%E8%AF%95%E9%99%A2%2C%E5%9B%BD%E7%BD%91%E5%B9%B3%E6%98%8C%E5%8E%BF%E4%BE%9B%E7%94%B5%E5%85%AC%E5%8F%B8%2C%E5%B9%B3%E6%98%8C%E5%8E%BF%E4%BA%BA%E6%B0%91%E6%94%BF%E5%BA%9C%E5%8A%9E%E5%85%AC%E5%AE%A4%2C%E4%BA%BA%E6%B0%91%E7%BD%91%2C%E5%B9%B3%E6%98%8C%E5%8E%BF%E5%8F%B8%E6%B3%95%E5%B1%80%2C%E5%B9%B3%E6%98%8C%E5%8E%BF%E7%BB%BC%E5%90%88%E8%A1%8C%E6%94%BF%E6%89%A7%E6%B3%95%E5%B1%80',
        'authenticatecenterjsessionid': jid,
        'bz_govc_SHIROJSESSIONID': bz_gov_id,
    }

    headers = {
        'Accept': '*/*',
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6',
        'Connection': 'keep-alive',
        'Origin': 'http://10.15.3.133:'+conf.jiyuehua_port,
        'Referer': 'http://10.15.3.133:83/fileCenter/uploadPage?uploadType=1&s=0.1754363371550416&isOpen=true&_=1722305351088',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/128.0.0.0 Safari/537.36 Edg/128.0.0.0',
    }

    # 获取当前时间并设置为中国标准时间（CST）
    current_time = datetime.now(pytz.timezone('Asia/Shanghai'))
    formatted_time = current_time.strftime('%a %b %d %Y %H:%M:%S GMT%z (中国标准时间)')

    script_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(script_dir, 'downloadfile',file_name)

    files = {
        'siteId': (None, conf.jiyuehua_siteid),
        'sessionId': (None, bz_gov_id),
        'type': (None, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
        'title': (None, parent_t),
        'contentId': (None, content_id),
        'filePath': (None, path_excel),
        'fileName': (None, article_t),
        'id': (None, 'WU_FILE_0'),
        'name': (None, file_name),
        'lastModifiedDate': (None, formatted_time),
        # 'size': (None, '43008'),
        'Filedata': (file_name, open(file_path, 'rb'), 'application/vnd.ms-excel'),
    }

    response = requests.post(
        'http://10.15.3.133:'+conf.jiyuehua_port+'/content/updateFileByPath',
        cookies=cookies,
        headers=headers,
        files=files,
        verify=False,
    )

    print(response.status_code)
    print(response.text)
    return response.json()




def get_libreoffice_path():
    """自动获取各平台的LibreOffice路径"""
    if sys.platform == 'win32':
        # Windows默认安装路径（可能需要根据实际安装位置调整）
        paths = [
            r"E:\soft3\libreoffice\program\soffice.exe"
        ]
    else:  # Linux/macOS
        paths = ["/usr/bin/libreoffice"]

    for path in paths:
        if os.path.exists(path):
            return path
    raise FileNotFoundError("未找到LibreOffice，请确认已安装并添加到系统路径")


def doc_to_docx(file_path):
    """跨平台doc转docx（使用LibreOffice）"""
    # 验证输入文件
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"输入文件不存在: {file_path}")

    # 准备输出路径
    output_dir = os.path.dirname(file_path)
    output_path = Path(file_path).with_suffix('.docx')

    # 构建命令行参数
    command = [
        get_libreoffice_path(),
        "--headless",
        "--convert-to", "docx",
        "--outdir", output_dir,
        file_path
    ]

    try:
        # 执行转换（超时60秒）
        result = subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError("文档转换超时，请检查文件是否损坏")

    # 检查转换结果
    if result.returncode != 0:
        error_msg = result.stderr.decode('utf-8', errors='ignore')
        raise RuntimeError(f"转换失败: {error_msg}")

    if not output_path.exists():
        available_files = "\n".join(os.listdir(output_dir))
        raise FileNotFoundError(
            f"输出文件未生成，目录中存在：\n{available_files}"
        )

    return str(output_path)

def find_and_replace_text(input_file, output_file, old_text, new_text):
    doc = fitz.open(input_file)

    # 正确获取PDF版本的方式
    pdf_version = doc.metadata.get("format", "未知版本").split()[-1]
    print(f"PDF规范版本：ISO 32000-1:{pdf_version} (对应PDF {doc.metadata['format']})")

    for page in doc:
        # 文本存在性验证（强化版）
        text_blocks = page.get_text("blocks")
        for i, block in enumerate(text_blocks):
            print(f"块 {i} 详细信息:")
            print(f"坐标范围：({block[0]}, {block[1]}) → ({block[2]}, {block[3]})")
            print(f"内容：{block[4]}")
            print(f"块类型：{'文本' if block[6] == 0 else '非文本'}\n{'-' * 40}")
        found = any(old_text in block[4] for block in text_blocks if block[6] == 0)
        print(f"文本存在性二次验证：{'找到' if found else '未找到'}")

        text_instances = page.search_for(old_text)
        print(f"原始匹配结果：{len(text_instances)}处")

        for rect in text_instances:
            # 添加调试标记
            red_rect = page.add_redact_annot(rect, fill=(1, 1, 0.8))  # 黄色背景
            page.apply_redactions()

            # 强制使用内置字体
            page.insert_text(
                rect.bl + (0, -rect.height * 0.1),  # 基线校准
                new_text,
                fontsize=rect.height * 0.8,
                fontname="china-s",  # 重要修正：指定内置中文字体
                color=(1, 0, 0),
                overlay=True
            )
            print(f"已写入：{new_text} | 坐标：{tuple(rect)}")

    # 版本兼容的保存方式
    doc.save(output_file, garbage=4, deflate=True, clean=True)
    print(f"最终文件大小：{os.path.getsize(output_file) // 1024}KB")


# 环境检查（必须运行）
print(f"PyMuPDF版本：{fitz.version}")  # 正确版本获取方式

import fitz
from itertools import product


def split_all_combinations(word, max_split=3):
    """生成所有可能的分割组合，最多拆分为3部分"""
    return [(word[:i], word[i:]) for i in range(1, len(word)) if i <= max_split or (len(word) - i) <= max_split]


def find_crossline_instances(page, old_text):
    text_blocks = [b for b in page.get_text("blocks") if b[6] == 0]
    matches = []

    # 遍历所有可能的分割组合
    for prefix, suffix in split_all_combinations(old_text):
        for i in range(len(text_blocks) - 1):
            curr_block = text_blocks[i][4].replace('\n', '')
            next_block = text_blocks[i + 1][4].replace('\n', '')

            # 检查跨块匹配
            if curr_block.endswith(prefix) and next_block.startswith(suffix):
                # 计算合并区域
                rect1 = fitz.Rect(text_blocks[i][:4])
                rect2 = fitz.Rect(text_blocks[i + 1][:4])
                combined_rect = rect1 | rect2

                matches.append({
                    "prefix_rect": rect1,
                    "suffix_rect": rect2,
                    "combined_rect": combined_rect,
                    "split_point": len(prefix)
                })

    return matches


def find_and_replace_text2(input_file, output_file, old_text, new_text):
    doc = fitz.open(input_file)

    for page in doc:
        # 首先处理单行匹配
        single_line_matches = page.search_for(old_text)
        for rect in single_line_matches:
            page.add_redact_annot(rect)
            page.apply_redactions()
            page.insert_text(
                rect.bl + (0, -rect.height * 0.1),
                new_text,
                fontsize=rect.height * 0.8,
                fontname="china-s",
                color=(1, 0, 0),
                overlay=True
            )

        # 处理跨行匹配
        cross_matches = find_crossline_instances(page, old_text)
        for match in cross_matches:
            # 创建跨行区域的红色删除线
            page.add_redact_annot(match["prefix_rect"])
            page.add_redact_annot(match["suffix_rect"])
            page.apply_redactions()

            # 在合并区域居中插入新文本
            center_point = match["combined_rect"].bl + (0, -match["combined_rect"].height * 0.2)
            page.insert_text(
                center_point,
                new_text,
                fontsize=match["combined_rect"].height * 0.7,
                fontname="china-s",
                color=(1, 0, 0),
                overlay=True
            )

    doc.save(output_file, garbage=4, deflate=True, clean=True)


if __name__ == "__main__":
    input_file = r"G:\project\python\pcxzf\pcxzf\bazhouqu\kaipumodify\modifyfile\downloadfile\20180904150156-752201\20180904150156-752201_101.pdf"
    output_file = r"G:\project\python\pcxzf\pcxzf\bazhouqu\kaipumodify\modifyfile\downloadfile\20180904150156-752201\modified_output.pdf"

    # 示例替换参数（需要您自行修改）
    old_text = "中华人民共和国国家发展改革委员会"
    new_text = "中华人民共和国国家发展和改革委员会"

    find_and_replace_text2(input_file, output_file, old_text, new_text)
