# -*- coding: utf-8 -*-
import os

from datetime import datetime
from nanjiang.kaipumodify.cfg import conf
import pytz
import requests
import xlrd
from openpyxl import load_workbook
from xlutils.copy import copy
from docx import Document
import pypandoc
import subprocess

import sys
from pathlib import Path

def download_file(url, local_filename):
    response = requests.get(url, stream=True, verify=False)
    response.raise_for_status()

    script_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(script_dir, 'downloadfile',local_filename)

    with open(file_path, 'wb') as file:
        for chunk in response.iter_content(chunk_size=8192):
            file.write(chunk)

    print(f"文件已下载并保存为: {file_path}")
    # time.sleep(5)
    return file_path



def modify_file_xls(file_path, unique_replacements_list):
    # 打开xls文件
    workbook_xls = xlrd.open_workbook(file_path, formatting_info=True)
    workbook_copy = copy(workbook_xls)


    # 遍历所有的sheet
    for sheet_index in range(workbook_xls.nsheets):
        sheet_xls = workbook_xls.sheet_by_index(sheet_index)
        sheet_copy = workbook_copy.get_sheet(sheet_index)

        # 遍历所有的单元格
        for row in range(sheet_xls.nrows):
            for col in range(sheet_xls.ncols):
                cell_value = sheet_xls.cell_value(row, col)

                for i in unique_replacements_list:

                    if isinstance(cell_value, str) and i["sensitiveWords"] in cell_value:
                        new_value = cell_value.replace(i["sensitiveWords"], i["recommendUpdate"])
                        sheet_copy.write(row, col, new_value)
                        print(
                            f"Sheet '{sheet_xls.name}' 中的单元格 '{row + 1},{col + 1}' 的值已从 '{cell_value}' 替换为 '{new_value}'")

    # 保存修改后的xls文件
    workbook_copy.save(file_path)
    print(f"文件 '{file_path}' 已成功修改并保存")


def modify_file_xlsx(file_path, unique_replacements_list):
    # 加载工作簿
    workbook = load_workbook(file_path)

    # 遍历所有的 sheet
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]

        # 遍历所有的单元格
        for row in sheet.iter_rows():
            for cell in row:
                for i in unique_replacements_list:
                    if isinstance(cell.value, str) and i["sensitiveWords"] in cell.value:
                        new_value = cell.value.replace(i["sensitiveWords"], i["recommendUpdate"])
                        cell.value = new_value
                        print(f"Sheet '{sheet_name}' 中的单元格 '{cell.coordinate}' 的值已从 '{i['sensitiveWords']}' 替换为 '{i['recommendUpdate']}'")

    # 保存修改后的工作簿
    workbook.save(file_path)
    print(f"文件 '{file_path}' 已成功修改并保存")


def modify_file_docx(file_path, unique_replacements_list):
    # 加载文档
    doc = Document(file_path)

    # 遍历所有段落
    for para in doc.paragraphs:
        for i in unique_replacements_list:
            if i["sensitiveWords"] in para.text:
                new_text = para.text.replace(i["sensitiveWords"], i["recommendUpdate"])
                para.text = new_text
                print(f"段落中的文本已从 '{i['sensitiveWords']}' 替换为 '{i['recommendUpdate']}'")

    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for i in unique_replacements_list:
                    if i["sensitiveWords"] in cell.text:
                        new_text = cell.text.replace(i["sensitiveWords"], i["recommendUpdate"])
                        cell.text = new_text
                        print(f"表格中的单元格文本已从 '{i['sensitiveWords']}' 替换为 '{i['recommendUpdate']}'")

    # 保存修改后的文档
    doc.save(file_path)
    print(f"文件 '{file_path}' 已成功修改并保存")




def modify_file(filename, item):
    script_dir = os.path.dirname(os.path.abspath(__file__))

    file_path = os.path.join(script_dir, 'downloadfile',filename)

    _, file_extension = os.path.splitext(file_path)

    unique_replacements = {}
    for i in item:
        key = (i['sensitiveWords'], i['recommendUpdate'])
        if key not in unique_replacements:
            unique_replacements[key] = i

    # 将字典转换为列表
    unique_replacements_list = list(unique_replacements.values())

    if file_extension == '.xls':
        modify_file_xls(file_path, unique_replacements_list)
    elif file_extension == '.xlsx':
        modify_file_xlsx(file_path, unique_replacements_list)
    elif  file_extension == '.docx':
        modify_file_docx(file_path, unique_replacements_list)
    elif file_extension == '.doc' :
        modify_file_doc(file_path, unique_replacements_list)
        # return 1
    return 0

def modify_file_doc(file_path_old, unique_replacements_list):
    try:
        output = doc_to_docx(file_path_old)
        print(f"转换成功！文件保存至：{output}")
        print(f"文件大小：{os.path.getsize(output) / 1024:.2f} KB")
    except Exception as e:
        print(f"错误：{str(e)}")
    # file_path_new = file_path_old.replace('.doc', '.docx')
    # pypandoc.convert_file(file_path_old, 'docx', outputfile=file_path)
    modify_file_docx(output, unique_replacements_list)





def uploadfile(jid, bz_gov_id, file_name, path_excel, parent_t, article_t, content_id):
    cookies = {
        'historyCookie': '%E5%8E%BF%E4%BA%BA%E6%B0%91%E6%94%BF%E5%BA%9C%E5%8A%9E%E5%85%AC%E5%AE%A4%2C%E4%B8%AD%E5%8D%8E%E4%BA%BA%E6%B0%91%E5%85%B1%E5%92%8C%E5%9B%BD%E6%95%99%E8%82%B2%E9%83%A8%2C%E4%BA%BA%E6%B0%91%E6%97%A5%E6%8A%A5%2C%E5%8E%BF%E8%87%AA%E7%84%B6%E8%B5%84%E6%BA%90%E5%92%8C%E8%A7%84%E5%88%92%E5%B1%80%2C%E5%9B%9B%E5%B7%9D%E7%9C%81%E6%95%99%E8%82%B2%E8%80%83%E8%AF%95%E9%99%A2%2C%E5%9B%BD%E7%BD%91%E5%B9%B3%E6%98%8C%E5%8E%BF%E4%BE%9B%E7%94%B5%E5%85%AC%E5%8F%B8%2C%E5%B9%B3%E6%98%8C%E5%8E%BF%E4%BA%BA%E6%B0%91%E6%94%BF%E5%BA%9C%E5%8A%9E%E5%85%AC%E5%AE%A4%2C%E4%BA%BA%E6%B0%91%E7%BD%91%2C%E5%B9%B3%E6%98%8C%E5%8E%BF%E5%8F%B8%E6%B3%95%E5%B1%80%2C%E5%B9%B3%E6%98%8C%E5%8E%BF%E7%BB%BC%E5%90%88%E8%A1%8C%E6%94%BF%E6%89%A7%E6%B3%95%E5%B1%80',
        'authenticatecenterjsessionid': jid,
        'bz_govc_SHIROJSESSIONID': bz_gov_id,
    }

    headers = {
        'Accept': '*/*',
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6',
        'Connection': 'keep-alive',
        'Origin': 'http://10.15.3.133:'+conf.jiyuehua_port,
        'Referer': 'http://10.15.3.133:83/fileCenter/uploadPage?uploadType=1&s=0.1754363371550416&isOpen=true&_=1722305351088',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/128.0.0.0 Safari/537.36 Edg/128.0.0.0',
    }

    # 获取当前时间并设置为中国标准时间（CST）
    current_time = datetime.now(pytz.timezone('Asia/Shanghai'))
    formatted_time = current_time.strftime('%a %b %d %Y %H:%M:%S GMT%z (中国标准时间)')

    script_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(script_dir, 'downloadfile',file_name)

    files = {
        'siteId': (None, conf.jiyuehua_siteid),
        'sessionId': (None, bz_gov_id),
        'type': (None, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
        'title': (None, parent_t),
        'contentId': (None, content_id),
        'filePath': (None, path_excel),
        'fileName': (None, article_t),
        'id': (None, 'WU_FILE_0'),
        'name': (None, file_name),
        'lastModifiedDate': (None, formatted_time),
        # 'size': (None, '43008'),
        'Filedata': (file_name, open(file_path, 'rb'), 'application/vnd.ms-excel'),
    }

    response = requests.post(
        'http://10.15.3.133:'+conf.jiyuehua_port+'/content/updateFileByPath',
        cookies=cookies,
        headers=headers,
        files=files,
        verify=False,
    )

    print(response.status_code)
    print(response.text)
    return response.json()




def get_libreoffice_path():
    """自动获取各平台的LibreOffice路径"""
    if sys.platform == 'win32':
        # Windows默认安装路径（可能需要根据实际安装位置调整）
        paths = [
            r"E:\soft3\libreoffice\program\soffice.exe"
        ]
    else:  # Linux/macOS
        paths = ["/usr/bin/libreoffice"]

    for path in paths:
        if os.path.exists(path):
            return path
    raise FileNotFoundError("未找到LibreOffice，请确认已安装并添加到系统路径")


def doc_to_docx(file_path):
    """跨平台doc转docx（使用LibreOffice）"""
    # 验证输入文件
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"输入文件不存在: {file_path}")

    # 准备输出路径
    output_dir = os.path.dirname(file_path)
    output_path = Path(file_path).with_suffix('.docx')

    # 构建命令行参数
    command = [
        get_libreoffice_path(),
        "--headless",
        "--convert-to", "docx",
        "--outdir", output_dir,
        file_path
    ]

    try:
        # 执行转换（超时60秒）
        result = subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError("文档转换超时，请检查文件是否损坏")

    # 检查转换结果
    if result.returncode != 0:
        error_msg = result.stderr.decode('utf-8', errors='ignore')
        raise RuntimeError(f"转换失败: {error_msg}")

    if not output_path.exists():
        available_files = "\n".join(os.listdir(output_dir))
        raise FileNotFoundError(
            f"输出文件未生成，目录中存在：\n{available_files}"
        )

    return str(output_path)




if __name__ == '__main__':

    # 测试您的文件路径（Windows环境）
    file_path = r'G:\project\python\pcxzf\pcxzf\bazhouqu\kaipumodify\modifyfile\downloadfile\rBUtIWeQhPaAHVkEAAXkAGN_bdk908.doc'

    try:
        output = doc_to_docx(file_path)
        print(f"转换成功！文件保存至：{output}")
        print(f"文件大小：{os.path.getsize(output) / 1024:.2f} KB")
    except Exception as e:
        print(f"错误：{str(e)}")

    docx_path = modify_file_doc(output,None)
    print(f"文件已转换为: {docx_path}")



